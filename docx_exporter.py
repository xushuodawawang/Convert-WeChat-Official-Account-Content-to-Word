from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from parser import ArticleData, ContentBlock, TextRun


@dataclass
class ExportOptions:
    font_name: str = "宋体"
    font_size: int = 12
    line_spacing: float = 1.5
    keep_images: bool = True
    link_position: str = "start"


class DocxExportError(Exception):
    pass


def export_article_to_docx_bytes(article: ArticleData, options: ExportOptions) -> bytes:
    try:
        document = Document()
        _configure_document_styles(document, options)
        _write_article(document, article, options)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except Exception as exc:  # pragma: no cover
        raise DocxExportError(f"Word 导出失败: {exc}") from exc


def _configure_document_styles(document: Document, options: ExportOptions) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = options.font_name
    normal_style.font.size = Pt(options.font_size)
    _set_east_asia_font(normal_style._element, options.font_name)

    _set_style_font(document.styles["Title"], options.font_name, max(options.font_size + 6, 18))
    _set_style_font(document.styles["Heading 1"], options.font_name, max(options.font_size + 2, 14))
    _set_style_font(document.styles["Heading 2"], options.font_name, max(options.font_size + 1, 13))
    _set_style_font(document.styles["Heading 3"], options.font_name, max(options.font_size, 12))
    _set_style_font(document.styles["Heading 4"], options.font_name, max(options.font_size, 12))


def _set_style_font(style, font_name: str, font_size: int) -> None:
    style.font.name = font_name
    style.font.size = Pt(font_size)
    _set_east_asia_font(style._element, font_name)


def _write_article(document: Document, article: ArticleData, options: ExportOptions) -> None:
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(article.title or "未命名文章")
    _apply_run_font(title_run, options.font_name, max(options.font_size + 6, 18), bold=True)
    _apply_line_spacing(title_paragraph, options.line_spacing)

    if options.link_position == "start":
        _write_metadata(document, article, options, include_link=True)
    else:
        _write_metadata(document, article, options, include_link=False)

    for block in article.blocks:
        _write_block(document, block, options)

    if options.link_position != "start":
        _write_link_footer(document, article.url, options)


def _write_metadata(
    document: Document,
    article: ArticleData,
    options: ExportOptions,
    include_link: bool,
) -> None:
    metadata_lines = [
        f"作者: {article.author or '未知'}",
        f"公众号: {article.publisher or '未知'}",
        f"发布时间: {article.publish_time or '未知'}",
        f"抓取方式: {article.source}",
    ]
    if include_link:
        metadata_lines.append(f"原文链接: {article.url}")

    for line in metadata_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        _apply_run_font(run, options.font_name, options.font_size)
        _apply_line_spacing(paragraph, options.line_spacing)

    document.add_paragraph()


def _write_link_footer(document: Document, url: str, options: ExportOptions) -> None:
    document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"原文链接: {url}")
    _apply_run_font(run, options.font_name, options.font_size)
    _apply_line_spacing(paragraph, options.line_spacing)


def _write_block(document: Document, block: ContentBlock, options: ExportOptions) -> None:
    if block.type == "heading":
        level = min(max(block.level, 1), 4)
        paragraph = document.add_paragraph(style=f"Heading {level}")
        _write_runs(paragraph, block.runs or [TextRun(block.text, bold=True)], options)
        _apply_line_spacing(paragraph, options.line_spacing)
        return

    if block.type == "image":
        if not options.keep_images or not block.image_bytes:
            if block.alt_text:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"[图片] {block.alt_text}")
                _apply_run_font(run, options.font_name, options.font_size, italic=True)
                _apply_line_spacing(paragraph, options.line_spacing)
            return

        paragraph = document.add_paragraph()
        _apply_line_spacing(paragraph, options.line_spacing)
        run = paragraph.add_run()
        image_stream = io.BytesIO(block.image_bytes)
        try:
            run.add_picture(image_stream, width=Inches(5.8))
        except Exception:
            fallback = paragraph.add_run(f"[图片插入失败] {block.alt_text or block.image_url or '未命名图片'}")
            _apply_run_font(fallback, options.font_name, options.font_size, italic=True)
        if block.alt_text:
            caption = document.add_paragraph()
            caption_run = caption.add_run(block.alt_text)
            _apply_run_font(caption_run, options.font_name, max(options.font_size - 1, 8), italic=True)
            _apply_line_spacing(caption, options.line_spacing)
        return

    paragraph = document.add_paragraph()
    _write_runs(paragraph, block.runs or [TextRun(block.text)], options)
    _apply_line_spacing(paragraph, options.line_spacing)


def _write_runs(paragraph, runs: list[TextRun], options: ExportOptions) -> None:
    if not runs:
        return

    for text_run in runs:
        if not text_run.text:
            continue
        run = paragraph.add_run(text_run.text)
        _apply_run_font(
            run,
            options.font_name,
            options.font_size,
            bold=text_run.bold,
            italic=text_run.italic,
        )


def _apply_run_font(run, font_name: str, font_size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    _set_east_asia_font(run._element, font_name)


def _apply_line_spacing(paragraph, line_spacing: float) -> None:
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(6)


def _set_east_asia_font(element, font_name: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
